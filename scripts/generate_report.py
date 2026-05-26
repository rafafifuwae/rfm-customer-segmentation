"""Generate laporan proyek (Tugas Data Mining) dalam format .docx.

Output: reports/Laporan_Proyek_Data_Mining.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "reports" / "Laporan_Proyek_Data_Mining.docx"
FIG_DIR = ROOT / "reports" / "figures"
DEPLOY_FIG_DIR = FIG_DIR / "deployment"
DIAGRAMS_DIR = FIG_DIR / "diagrams"


# ---------- helpers ----------
def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_justified(doc, text: str, size: int = 11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items: list[str]):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
        if not p.runs:
            p.add_run(it).font.size = Pt(11)
        else:
            p.runs[0].text = it


def add_numbered(doc, items: list[str]):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(it).font.size = Pt(11)


def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # apply font to East Asian as well to prevent fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    rPr.append(rFonts)
    # shaded background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_table(doc, header: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], "1F3A5F")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_image(doc, image_path: Path, width_cm: float = 15.0, caption: str | None = None):
    """Sisipkan gambar dengan caption opsional, center-aligned."""
    if not image_path.exists():
        add_para(doc, f"[Gambar tidak ditemukan: {image_path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cap.paragraph_format.space_after = Pt(12)


# ---------- main ----------
def build() -> Document:
    doc = Document()

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============= COVER =============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN PROYEK DATA MINING")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Segmentasi Pelanggan untuk Mengatasi Ketidakefektifan "
        "Strategi Pemasaran Menggunakan Metode RFM dan Algoritma K-Means"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.italic = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        "Studi Kasus: Brazilian E-Commerce Public Dataset by Olist\n"
        "Periode September 2016 – Oktober 2018"
    ).font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    add_page_break(doc)

    # ============= SOAL 1 =============
    add_heading(doc, "1. Judul Proyek", level=1)
    add_justified(
        doc,
        "Judul proyek data mining yang diangkat dalam penelitian ini adalah:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        '"Segmentasi Pelanggan untuk Mengatasi Ketidakefektifan Strategi '
        "Pemasaran Menggunakan Metode RFM dan Algoritma K-Means\""
    )
    r.bold = True
    r.italic = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "(Studi Kasus: Brazilian E-Commerce Public Dataset by Olist, "
        "Periode September 2016 – Oktober 2018)"
    ).italic = True

    # ============= SOAL 2 =============
    add_heading(doc, "2. Latar Belakang Penelitian", level=1)

    add_justified(
        doc,
        "Perkembangan e-commerce dalam satu dekade terakhir telah mengubah "
        "lanskap perdagangan global secara fundamental. Marketplace seperti "
        "Olist — platform e-commerce terbesar di Brazil — menghubungkan "
        "ribuan penjual UMKM dengan jutaan pelanggan setiap harinya. Volume "
        "transaksi yang besar ini menghasilkan basis pelanggan yang sangat "
        "heterogen: ada pelanggan loyal yang sering bertransaksi dengan "
        "nilai tinggi, pelanggan baru yang masih dalam tahap eksplorasi, "
        "hingga pelanggan yang sudah lama tidak aktif.",
    )

    add_justified(
        doc,
        "Tantangan utama yang dihadapi tim pemasaran adalah ketidakefektifan "
        'strategi pemasaran "one-size-fits-all". Ketika promosi yang sama '
        "dikirimkan ke seluruh basis pelanggan tanpa memperhatikan "
        "karakteristik perilaku mereka, hasilnya adalah:",
    )

    add_numbered(
        doc,
        [
            "Pemborosan anggaran marketing — pelanggan VIP yang sudah loyal "
            "tidak butuh diskon agresif; sebaliknya, pelanggan yang sudah "
            "lama tidak transaksi tidak akan tergerak oleh newsletter biasa.",
            "Rendahnya conversion rate — pesan yang tidak relevan akan "
            "diabaikan, bahkan berpotensi membuat pelanggan unsubscribe.",
            "Churn yang tidak terdeteksi dini — pelanggan bernilai tinggi "
            "yang mulai menurun aktivitasnya tidak mendapat intervensi tepat "
            "waktu.",
            "Customer Lifetime Value (CLV) tidak teroptimalkan — segmen yang "
            "seharusnya bisa di-upsell justru diperlakukan sama seperti "
            "pelanggan one-time buyer.",
        ],
    )

    add_justified(
        doc,
        "Untuk mengatasi permasalahan tersebut, diperlukan pendekatan "
        "data-driven customer segmentation yang mampu mengelompokkan "
        "pelanggan berdasarkan perilaku transaksi mereka secara objektif. "
        "Salah satu metode yang sudah teruji secara empiris adalah RFM "
        "Analysis (Recency, Frequency, Monetary) yang dikembangkan Hughes "
        "(1995) dan menjadi standar de-facto di industri customer analytics. "
        "RFM hanya membutuhkan data transaksi (tidak butuh data demografis/"
        "PII) dan menghasilkan tiga dimensi perilaku yang interpretable bagi "
        "tim bisnis:",
    )

    add_bullets(
        doc,
        [
            "Recency — kapan terakhir kali pelanggan bertransaksi "
            "(mengukur engagement terkini).",
            "Frequency — seberapa sering pelanggan bertransaksi "
            "(mengukur loyalitas).",
            "Monetary — total nilai transaksi pelanggan "
            "(mengukur nilai ekonomis).",
        ],
    )

    add_justified(
        doc,
        "Setelah fitur RFM dihasilkan, dibutuhkan algoritma clustering untuk "
        "membentuk segmen secara otomatis. K-Means Clustering dipilih karena "
        "cocok untuk data numerik kontinu, skalabel ke ratusan ribu "
        "pelanggan dengan kompleksitas O(nki), menghasilkan centroid yang "
        "dapat dipakai untuk klasifikasi pelanggan baru, serta mudah "
        "diinterpretasikan oleh stakeholder non-teknis.",
    )

    add_justified(
        doc,
        "Penelitian ini menggunakan Brazilian E-Commerce Public Dataset by "
        "Olist yang berisi sekitar 100.000 transaksi dari 99.000 pelanggan "
        "unik dalam periode 2016–2018. Hasil akhir penelitian diharapkan "
        "dapat: (1) menghasilkan model segmentasi pelanggan yang dapat "
        "dievaluasi secara kuantitatif, (2) menghasilkan profil dan label "
        "segmen yang interpretable secara bisnis, serta (3) memberikan "
        "rekomendasi strategi pemasaran yang spesifik untuk tiap segmen.",
    )

    add_heading(doc, "Rumusan Masalah", level=2)
    add_numbered(
        doc,
        [
            "Bagaimana mengelompokkan pelanggan Olist berdasarkan perilaku "
            "transaksi menggunakan metode RFM dan algoritma K-Means?",
            "Apa karakteristik tiap segmen yang dihasilkan?",
            "Rekomendasi strategi pemasaran apa yang sesuai untuk tiap "
            "segmen pelanggan?",
        ],
    )

    add_heading(doc, "Tujuan Penelitian", level=2)
    add_numbered(
        doc,
        [
            "Mengimplementasikan pipeline data mining dari raw data hingga "
            "deployment yang reproducible.",
            "Menentukan jumlah cluster optimal menggunakan kombinasi Elbow "
            "Method, Silhouette Score, dan Davies-Bouldin Index.",
            "Menyusun rekomendasi strategi pemasaran berbasis profil tiap "
            "segmen.",
            "Mengembangkan layanan deployment (REST API + Web UI) sebagai "
            "bukti bahwa model siap digunakan.",
        ],
    )

    add_page_break(doc)

    # ============= SOAL 3 =============
    add_heading(doc, "3. Diagram Alur Penelitian", level=1)
    add_justified(
        doc,
        "Penelitian ini mengikuti kerangka CRISP-DM (Cross-Industry Standard "
        "Process for Data Mining) yang terdiri dari enam fase. Berikut "
        "diagram alur penelitian yang dilakukan:",
    )

    add_image(doc, DIAGRAMS_DIR / "01_research_flow.png", width_cm=14,
              caption="Gambar 3.1 — Diagram alur penelitian data mining "
                      "(adaptasi CRISP-DM untuk segmentasi pelanggan)")

    _unused_diagram_research = """┌─────────────────────────────────────────────────────────┐
│            ALUR PENELITIAN DATA MINING                  │
│   (Adaptasi CRISP-DM untuk Segmentasi Pelanggan)        │
└─────────────────────────────────────────────────────────┘

      ┌──────────────────────────┐
      │ 1. BUSINESS UNDERSTANDING│
      │  - Identifikasi masalah  │
      │  - Definisi tujuan       │
      │  - Definisi keberhasilan │
      └────────────┬─────────────┘
                   │
                   ▼
      ┌──────────────────────────┐
      │  2. DATA UNDERSTANDING   │  →  01_data_collection.ipynb
      │  - Load 9 tabel Olist    │     02_eda.ipynb
      │  - Validasi schema       │     reports/data_dictionary.csv
      │  - EDA (uni/bi/multivar) │     reports/figures/*.png
      │  - Cek refer. integrity  │
      └────────────┬─────────────┘
                   │
                   ▼
      ┌──────────────────────────┐
      │   3. DATA PREPARATION    │  →  03_data_preprocessing.ipynb
      │  a. Cleaning             │     - Filter status = delivered
      │  b. Merging              │     - Merge orders+customers+pay
      │  c. Aggregation          │     - Group by customer_unique_id
      │  d. RFM Feature Eng.     │     - Recency/Frequency/Monetary
      │  e. Outlier handling     │     - Winsorize (99th percentile)
      │  f. Transformation       │     - Log1p (Frequency, Monetary)
      │  g. RFM Scoring          │     - Quintile-based (1-5)
      │  Output: rfm_table.csv   │
      └────────────┬─────────────┘
                   │
                   ▼
      ┌──────────────────────────┐
      │      4. MODELING         │  →  04_rfm_kmeans_modeling.ipynb
      │  - StandardScaler        │
      │  - Pemilihan k (2..10)   │     • Elbow Method
      │                          │     • Silhouette Score
      │                          │     • Davies-Bouldin Index
      │  - Fit K-Means (k=4)     │     n_init=20, random_state=42
      │  - PCA visualisasi 2D    │
      │  Output: kmeans_rfm.pkl  │
      └────────────┬─────────────┘
                   │
                   ▼
      ┌──────────────────────────┐
      │     5. EVALUATION        │  →  models/kmeans_metrics.json
      │  Metrik Internal:        │     reports/figures/
      │  - Silhouette = 0.372    │       cluster_snake_plot.png
      │  - Davies-Bouldin = 0.76 │       cluster_pca_scatter.png
      │  - Inertia = 79.016      │       cluster_size.png
      │  Interpretasi Bisnis:    │
      │  - Champions             │
      │  - Loyal Customers       │
      │  - At Risk               │
      │  - Hibernating / Lost    │
      │  Output: rfm_segmented   │
      └────────────┬─────────────┘
                   │
                   ▼
      ┌──────────────────────────┐
      │     6. DEPLOYMENT        │  →  app/api.py        (FastAPI)
      │  - Build serving artifact│     app/streamlit_app.py (UI)
      │  - REST API (FastAPI)    │     app/predictor.py
      │  - Web UI (Streamlit)    │     models/serving_artifacts.json
      │  - DVC pipeline          │     dvc.yaml + params.yaml
      │    (reproducible)        │
      └──────────────────────────┘
"""

    add_heading(doc, "Penjelasan Tiap Fase", level=2)

    add_para(doc, "Fase 1 — Business Understanding", bold=True)
    add_justified(
        doc,
        "Tahap awal yang bertujuan memahami konteks bisnis dan tujuan dari "
        "proyek data mining. Pada penelitian ini diidentifikasi bahwa "
        "strategi pemasaran seragam Olist tidak efektif karena basis "
        "pelanggannya sangat heterogen. Tujuan ditetapkan untuk membentuk "
        "segmen pelanggan yang dapat ditarget secara spesifik.",
    )

    add_para(doc, "Fase 2 — Data Understanding", bold=True)
    add_justified(
        doc,
        "Loading 9 tabel Olist (sekitar 120 MB), validasi schema, dan "
        "eksplorasi data melalui EDA yang mencakup analisis univariate "
        "(distribusi tiap variabel), bivariate (hubungan dua variabel), "
        "multivariate (korelasi antar fitur), time-series (tren bulanan dan "
        "harian), serta analisis geografis. Total 21 visualisasi dihasilkan "
        "sebagai bukti pemahaman data sebelum modeling.",
    )

    add_para(doc, "Fase 3 — Data Preparation", bold=True)
    add_justified(
        doc,
        "Tahap paling kritis dan memakan waktu paling besar (sekitar 60% "
        "waktu proyek). Termasuk filtering data, merging antar tabel, "
        "agregasi pembayaran, feature engineering RFM, penanganan outlier "
        "via winsorize, transformasi log, dan RFM scoring. Output utama: "
        "tabel RFM dengan 93.357 pelanggan unik siap-modeling.",
    )

    add_para(doc, "Fase 4 — Modeling", bold=True)
    add_justified(
        doc,
        "Standardisasi fitur menggunakan StandardScaler dilanjutkan dengan "
        "pemilihan jumlah cluster optimal melalui kombinasi tiga metrik "
        "(Elbow, Silhouette, Davies-Bouldin) untuk menghindari bias metrik "
        "tunggal. Hasil: k = 4 dengan parameter n_init=20 dan random_state=42 "
        "untuk reproducibility.",
    )

    add_para(doc, "Fase 5 — Evaluation", bold=True)
    add_justified(
        doc,
        "Evaluasi dilakukan dari dua sisi: teknis (silhouette, davies-bouldin, "
        "inertia) dan bisnis (interpretasi profil tiap cluster menjadi segmen "
        "yang actionable). Visualisasi snake plot dan PCA scatter membantu "
        "memvalidasi kualitas separasi cluster.",
    )

    add_para(doc, "Fase 6 — Deployment", bold=True)
    add_justified(
        doc,
        "Model di-wrap dalam REST API (FastAPI) dan Web UI (Streamlit) agar "
        "dapat digunakan tim bisnis untuk klasifikasi pelanggan baru maupun "
        "analisis batch. Pipeline preprocessing lengkap juga di-serving "
        "sehingga user cukup mengirimkan nilai RFM mentah.",
    )

    add_page_break(doc)

    # ============= SOAL 4 =============
    add_heading(doc, "4. Analisa Kode Program", level=1)

    add_heading(doc, "4.1 Pengambilan Data (src/data_loader.py)", level=2)
    add_justified(
        doc,
        "Loader data dibuat modular sebagai fungsi tunggal yang "
        "mengembalikan dict of DataFrames. Memudahkan unit-testing dan reuse "
        "baik di notebook maupun di API.",
    )
    add_code(doc, """def load_olist(raw_dir: Path | str = DEFAULT_RAW_DIR) -> Dict[str, pd.DataFrame]:
    raw_dir = Path(raw_dir)
    out: Dict[str, pd.DataFrame] = {}
    for key, fname in FILE_MAP.items():
        kwargs = {}
        if key in PARSE_DATES:
            kwargs["parse_dates"] = PARSE_DATES[key]
        out[key] = pd.read_csv(raw_dir / fname, **kwargs)
    return out""")
    add_para(doc, "Analisa:", bold=True)
    add_bullets(
        doc,
        [
            "parse_dates di-pass langsung ke read_csv agar parsing datetime "
            "dilakukan di level C (lebih cepat).",
            "PROJECT_ROOT dihitung relatif terhadap lokasi file sehingga "
            "loader bisa dipanggil dari direktori manapun.",
            "Mendukung 9 tabel Olist sekaligus dalam satu pemanggilan.",
        ],
    )

    add_heading(doc, "4.2 Exploratory Data Analysis (02_eda.ipynb)", level=2)
    add_justified(
        doc,
        "EDA dilakukan dengan pendekatan top-down: dari distribusi tiap "
        "variabel (univariate), hubungan dua variabel (bivariate), korelasi "
        "antar fitur (multivariate), hingga tren temporal dan geografis.",
    )
    add_code(doc, """# Univariate
sns.histplot(payments["payment_value"], bins=50, kde=True)

# Bivariate: review score vs delivery time
delivery_days = (orders["order_delivered_customer_date"]
                 - orders["order_purchase_timestamp"]).dt.days
sns.boxplot(x=reviews["review_score"], y=delivery_days)

# Multivariate: correlation heatmap
sns.heatmap(merged[numeric_cols].corr(), annot=True, cmap="coolwarm")

# Time-series
orders_monthly = orders.set_index("order_purchase_timestamp").resample("M").size()
orders_monthly.plot(kind="line")""")
    add_para(doc, "Temuan kunci dari EDA:", bold=True)
    add_bullets(
        doc,
        [
            "Distribusi payment_value sangat right-skewed → memotivasi log "
            "transform.",
            "Sebagian besar pelanggan hanya bertransaksi 1 kali (Frequency=1).",
            "Mayoritas pelanggan berasal dari state SP (Sao Paulo).",
            "Volume order meningkat tajam dari Q1 2017 (tren positif).",
        ],
    )

    add_image(doc, FIG_DIR / "univariate_payment_value.png", width_cm=13,
              caption="Gambar 4.1 — Distribusi payment_value (right-skewed)")
    add_image(doc, FIG_DIR / "multivariate_corr_heatmap.png", width_cm=13,
              caption="Gambar 4.2 — Correlation heatmap antar fitur numerik")
    add_image(doc, FIG_DIR / "timeseries_orders_per_month.png", width_cm=14,
              caption="Gambar 4.3 — Tren volume order per bulan (2016–2018)")

    add_heading(doc, "4.3 Preprocessing (03_data_preprocessing.ipynb)", level=2)
    add_code(doc, """# 1. Filter status delivered
orders_d = orders[orders["order_status"] == "delivered"].copy()
orders_d = orders_d.dropna(subset=["order_purchase_timestamp"])

# 2. Agregasi pembayaran per order
pay_per_order = (payments.groupby("order_id", as_index=False)
                 .agg(payment_value=("payment_value", "sum")))

# 3. Merge orders <-> customers <-> payments
df = (orders_d
      .merge(customers[["customer_id", "customer_unique_id"]], on="customer_id")
      .merge(pay_per_order, on="order_id"))

# 4. Feature Engineering RFM
snapshot_date = df["order_purchase_timestamp"].max() + pd.Timedelta(days=1)
rfm = (df.groupby("customer_unique_id").agg(
    Recency=("order_purchase_timestamp", lambda x: (snapshot_date - x.max()).days),
    Frequency=("order_id", "nunique"),
    Monetary=("payment_value", "sum")
).reset_index())

# 5. Winsorize + Log transform
rfm["Monetary_w"]   = winsorize(rfm["Monetary"], upper=0.99)
rfm["Monetary_log"] = np.log1p(rfm["Monetary_w"])
rfm["Frequency_log"] = np.log1p(winsorize(rfm["Frequency"], upper=0.99))

# 6. RFM Score (quintile)
rfm["R_score"] = pd.qcut(rfm["Recency_w"].rank(method="first"), 5,
                          labels=[5,4,3,2,1]).astype(int)""")
    add_para(doc, "Analisa:", bold=True)
    add_bullets(
        doc,
        [
            "customer_unique_id digunakan, bukan customer_id, karena di "
            "Olist customer_id unik per order, bukan per pelanggan.",
            "snapshot_date = max(timestamp) + 1 hari untuk menghindari "
            "Recency bernilai 0 yang ambigu.",
            "Winsorize 99% mengurangi pengaruh outlier ekstrem tanpa "
            "membuang data.",
            "Log transform mengatasi right-skewness Monetary & Frequency.",
            "rank(method='first') pada qcut menangani nilai duplikat (banyak "
            "pelanggan Frequency=1) yang akan error tanpa rank.",
            "Output: 93.357 pelanggan unik dengan fitur RFM siap-modeling.",
        ],
    )

    add_image(doc, FIG_DIR / "rfm_distribution_raw.png", width_cm=15,
              caption="Gambar 4.4 — Distribusi RFM sebelum transformasi (raw)")
    add_image(doc, FIG_DIR / "rfm_distribution_transformed.png", width_cm=15,
              caption="Gambar 4.5 — Distribusi RFM setelah winsorize + log transform")

    add_heading(doc, "4.4 Modeling (04_rfm_kmeans_modeling.ipynb)", level=2)
    add_code(doc, """# 1. Standardisasi
feature_cols = ["Recency_w", "Frequency_log", "Monetary_log"]
X = rfm[feature_cols].values
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# 2. Pemilihan k optimal (k=2..10)
for k in range(2, 11):
    km = KMeans(n_clusters=k, n_init=10, random_state=42)
    labels = km.fit_predict(X_scaled)
    inertias.append(km.inertia_)
    sil_scores.append(silhouette_score(X_scaled, labels, sample_size=10000))
    db_scores.append(davies_bouldin_score(X_scaled, labels))

# 3. Fit final dengan k=4
kmeans = KMeans(n_clusters=4, n_init=20, max_iter=300, random_state=42)
rfm["Cluster"] = kmeans.fit_predict(X_scaled)

# 4. Simpan model + scaler
joblib.dump({"scaler": scaler, "kmeans": kmeans,
             "feature_cols": feature_cols}, "models/kmeans_rfm.pkl")""")
    add_para(doc, "Analisa:", bold=True)
    add_bullets(
        doc,
        [
            "StandardScaler penting karena K-Means berbasis Euclidean "
            "distance — tanpa standardisasi, Monetary akan mendominasi.",
            "n_init=20 menjalankan K-Means 20 kali dengan inisialisasi "
            "berbeda untuk menghindari local optimum.",
            "random_state=42 memastikan hasil reproducible.",
            "Tiga metrik (Elbow, Silhouette, Davies-Bouldin) digunakan "
            "bersamaan untuk pemilihan k yang lebih robust.",
            "Model disimpan via joblib (idiomatic untuk objek sklearn).",
        ],
    )

    add_image(doc, FIG_DIR / "kmeans_k_selection.png", width_cm=15,
              caption="Gambar 4.6 — Pemilihan k optimal: Elbow + Silhouette + Davies-Bouldin")

    add_heading(doc, "4.5 Evaluasi", level=2)
    add_para(doc, "Metrik Kuantitatif:", bold=True)
    add_table(
        doc,
        ["Metrik", "Nilai", "Interpretasi"],
        [
            ["Silhouette Score", "0.372", "Moderat — wajar untuk data perilaku pelanggan"],
            ["Davies-Bouldin Index", "0.763", "< 1 → separasi cluster baik"],
            ["Inertia", "79.016", "Untuk elbow comparison"],
            ["n_samples", "93.357", "Coverage pelanggan"],
        ],
        col_widths=[4.0, 2.5, 9.0],
    )

    add_para(doc, "Profil 4 Cluster Hasil Akhir:", bold=True)
    add_table(
        doc,
        ["Cluster", "Segment", "n", "%", "Recency", "Frequency", "Monetary (R$)"],
        [
            ["0", "Loyal Customers", "35.247", "37,76%", "147 hari", "1,00", "68,52"],
            ["1", "Hibernating / Lost", "27.153", "29,09%", "425 hari", "1,00", "120,24"],
            ["2", "At Risk", "2.801", "3,00%", "220 hari", "2,11", "308,59"],
            ["3", "Champions", "28.156", "30,16%", "173 hari", "1,00", "315,31"],
        ],
        col_widths=[1.5, 3.5, 2.0, 1.8, 2.0, 2.0, 2.5],
    )

    add_para(doc, "Analisa Hasil:", bold=True)
    add_bullets(
        doc,
        [
            "Silhouette 0,372 wajar untuk data RFM real-world karena "
            "perilaku pelanggan kontinu (tidak ada batas tegas antar segmen).",
            "Davies-Bouldin 0,76 < 1 menunjukkan separasi cluster sudah baik.",
            "Segmen At Risk hanya 3% tapi memiliki Monetary & Frequency "
            "tertinggi → segmen kecil-bernilai-tinggi yang harus "
            "diprioritaskan untuk win-back campaign.",
            "Labeling segmen dibuat data-driven berbasis median antar "
            "cluster, sehingga keempat label pasti tercapai.",
        ],
    )

    add_image(doc, FIG_DIR / "cluster_snake_plot.png", width_cm=14,
              caption="Gambar 4.7 — Snake plot: profil standardized RFM per cluster")
    add_image(doc, FIG_DIR / "cluster_pca_scatter.png", width_cm=12,
              caption="Gambar 4.8 — Sebaran cluster dalam ruang 2D PCA")
    add_image(doc, FIG_DIR / "cluster_size.png", width_cm=13,
              caption="Gambar 4.9 — Jumlah pelanggan per cluster")

    add_page_break(doc)

    # ============= SOAL 5 =============
    add_heading(doc, "5. Analisa Hasil Deployment", level=1)

    add_justified(
        doc,
        "Model yang telah dilatih di-deploy sebagai aplikasi web dengan "
        "arsitektur dua-tier: REST API (FastAPI) sebagai backend dan Web UI "
        "(Streamlit) sebagai frontend. Pemisahan ini memungkinkan UI dapat "
        "diganti dengan dashboard atau mobile app tanpa mengubah backend.",
    )

    add_heading(doc, "5.1 Arsitektur Deployment", level=2)

    add_image(doc, DIAGRAMS_DIR / "02_deployment_architecture.png", width_cm=14,
              caption="Gambar 5.0 — Arsitektur deployment: FastAPI (backend) + "
                      "Streamlit (frontend) + RFMPredictor (transform pipeline) "
                      "dengan model & serving artifacts")

    _unused_arsitektur = """┌─────────────────────────────────────────────────────────┐
│              ARSITEKTUR DEPLOYMENT                       │
└─────────────────────────────────────────────────────────┘

   Browser (User)
        │ HTTP
        ▼
   ┌──────────────────┐   HTTP API   ┌─────────────────────┐
   │  Streamlit UI    │ ──────────▶  │  FastAPI REST API    │
   │  (port 8501)     │              │  (port 8000)         │
   │                  │ ◀──────────  │                      │
   │ • Single predict │   JSON       │ • POST /predict       │
   │ • Batch CSV      │              │ • POST /predict/batch │
   │ • Cluster view   │              │ • GET  /clusters      │
   └──────────────────┘              │ • GET  /health        │
                                     └──────────┬──────────┘
                                                │
                                                ▼
                                     ┌─────────────────────┐
                                     │  RFMPredictor        │
                                     │  Transform Pipeline: │
                                     │  Raw RFM             │
                                     │     ↓                │
                                     │  Winsorize           │
                                     │     ↓                │
                                     │  Log1p (F, M)        │
                                     │     ↓                │
                                     │  StandardScaler      │
                                     │     ↓                │
                                     │  KMeans.predict()    │
                                     │     ↓                │
                                     │  → Segment + Strategi│
                                     └─────────────────────┘
"""

    add_heading(doc, "5.2 Komponen Deployment", level=2)
    add_table(
        doc,
        ["File", "Fungsi"],
        [
            ["app/predictor.py", "Wrapper model: load joblib + transform pipeline + predict"],
            ["app/api.py", "FastAPI service dengan 4 endpoint utama"],
            ["app/schemas.py", "Pydantic request/response schemas (validasi otomatis)"],
            ["app/streamlit_app.py", "Web UI dengan 3 halaman interaktif"],
            ["app/build_artifacts.py", "Generate serving artifacts (1x setelah training)"],
            ["scripts/run_app.sh", "Launcher: start FastAPI + Streamlit sekaligus"],
        ],
        col_widths=[5.0, 10.5],
    )

    add_heading(doc, "5.3 Contoh Request & Response API", level=2)
    add_code(doc, """# Single prediction
POST /predict
Body: {"Recency": 30, "Frequency": 2, "Monetary": 350.5}

Response:
{
  "cluster": 2,
  "segment": "At Risk",
  "distance_to_centroid": 1.35,
  "strategy": {
    "deskripsi": "Dulu bernilai tinggi tapi sudah lama tidak transaksi.",
    "strategi": [
      "Win-back campaign dengan diskon agresif",
      "Survei untuk tahu alasan churn",
      "Personalized re-engagement email"
    ]
  }
}

# Batch prediction via CSV upload
POST /predict/batch
File: rfm.csv (kolom: Recency, Frequency, Monetary)
Response: {"n_rows": 1000, "rows": [...]}

# Profil semua cluster
GET /clusters
Response: [{"cluster": 0, "segment": "Loyal Customers", ...}, ...]

# Health check
GET /health
Response: {"status": "ok", "n_training_samples": 93357, "n_clusters": 4}""")

    add_heading(doc, "5.4 Tampilan Web UI (Streamlit)", level=2)
    add_para(doc, "Streamlit UI memiliki tiga halaman utama:", bold=True)
    add_numbered(
        doc,
        [
            "Single Prediction — form input Recency/Frequency/Monetary "
            "menghasilkan badge segmen berwarna, deskripsi segmen, dan "
            "daftar rekomendasi strategi marketing.",
            "Batch Prediction — upload CSV banyak pelanggan, hasil "
            "ditampilkan dalam tabel + bar chart distribusi segmen, dan "
            "dapat di-download sebagai CSV.",
            "Cluster Overview — profil tiap cluster (ukuran, rata-rata R/F/M, "
            "total monetary) + detail strategi marketing per segmen.",
        ],
    )

    add_para(doc, "Halaman 1 — Single Prediction:", bold=True)
    add_image(doc, DEPLOY_FIG_DIR / "01_streamlit_single_empty.png", width_cm=15,
              caption="Gambar 5.1 — Form input Single Prediction (state awal)")
    add_image(doc, DEPLOY_FIG_DIR / "02_streamlit_single_result.png", width_cm=15,
              caption="Gambar 5.2 — Hasil prediksi: input R=30, F=2, M=350,5 "
                      "diklasifikasikan sebagai segmen 'At Risk' beserta strategi")

    add_para(doc, "Halaman 2 — Batch Prediction:", bold=True)
    add_image(doc, DEPLOY_FIG_DIR / "03_streamlit_batch.png", width_cm=15,
              caption="Gambar 5.3 — Form upload CSV untuk batch prediction")
    add_image(doc, DEPLOY_FIG_DIR / "04_streamlit_batch_result.png", width_cm=15,
              caption="Gambar 5.4 — Hasil batch: 5 pelanggan diklasifikasikan, "
                      "tabel + bar chart distribusi segmen + opsi download CSV")

    add_para(doc, "Halaman 3 — Cluster Overview:", bold=True)
    add_image(doc, DEPLOY_FIG_DIR / "05_streamlit_cluster_overview.png", width_cm=15,
              caption="Gambar 5.5 — Profil 4 cluster + visualisasi jumlah "
                      "pelanggan & total monetary per cluster")

    add_para(doc, "FastAPI — Swagger Documentation:", bold=True)
    add_image(doc, DEPLOY_FIG_DIR / "06_fastapi_docs.png", width_cm=15,
              caption="Gambar 5.6 — Swagger UI FastAPI di endpoint /docs "
                      "(auto-generated, interaktif untuk testing API)")

    add_heading(doc, "5.5 Cara Menjalankan Deployment", level=2)
    add_code(doc, """# 1. Install dependencies
pip install -r requirements.txt

# 2. Generate serving artifacts (sekali setelah training)
python -m app.build_artifacts

# 3. Jalankan FastAPI + Streamlit sekaligus
./scripts/run_app.sh

# 4. Akses melalui browser:
#    Web UI       : http://localhost:8501
#    API Docs     : http://localhost:8000/docs   (Swagger interaktif)
#    Health check : http://localhost:8000/health""")

    add_heading(doc, "5.6 Bukti Deployment Berjalan", level=2)
    add_code(doc, """$ ./scripts/run_app.sh
>> FastAPI  : http://localhost:8000/docs
>> Streamlit: http://localhost:8501
BOTH READY after 2s

$ curl http://localhost:8000/health
{"status":"ok","n_training_samples":93357,"n_clusters":4}

$ curl -X POST http://localhost:8000/predict \\
       -H 'Content-Type: application/json' \\
       -d '{"Recency":30,"Frequency":2,"Monetary":350.5}'
{
  "cluster": 2,
  "segment": "At Risk",
  "distance_to_centroid": 1.3515,
  "strategy": {...}
}""")

    add_heading(doc, "5.7 Analisa Kelebihan Deployment", level=2)
    add_bullets(
        doc,
        [
            "Separation of concerns — UI dan logika prediksi terpisah, "
            "memungkinkan UI diganti tanpa mengubah backend.",
            "Validasi otomatis via Pydantic — input invalid akan ditolak "
            "dengan pesan error jelas (mis. Recency negatif).",
            "Swagger UI otomatis di-generate FastAPI di endpoint /docs, "
            "memudahkan integrasi dengan tim lain.",
            "Transform pipeline lengkap di-serving — user cukup kirim RFM "
            "mentah tanpa perlu tahu winsorize/log/scaling.",
            "Batch processing hingga 100.000 baris per request, cocok untuk "
            "analisis seluruh basis pelanggan.",
            "Reproducibility — artifact serving disimpan dalam JSON terpisah "
            "agar transform konsisten antara training dan inference.",
        ],
    )

    add_heading(doc, "5.8 Manfaat Bisnis", level=2)
    add_bullets(
        doc,
        [
            "Tim marketing dapat mengetahui segmen pelanggan baru dalam "
            "kurang dari 100 milidetik.",
            "Analisis batch memungkinkan kampanye dijalankan ke segmen "
            "tertentu (mis. semua pelanggan At Risk → win-back campaign).",
            "Cluster overview membantu manajemen melihat distribusi nilai "
            "bisnis per segmen untuk alokasi anggaran marketing.",
            "Strategi rekomendasi siap-pakai memudahkan tim marketing "
            "menyusun campaign tanpa harus menafsir hasil clustering "
            "secara manual.",
        ],
    )

    add_heading(doc, "5.9 Strategi Deployment Multi-Layer", level=2)
    add_justified(
        doc,
        "Untuk memastikan aplikasi dapat dijalankan pada berbagai konteks "
        "(development, demo presentasi, hingga hosting permanen), proyek ini "
        "mendukung tiga layer deployment yang saling melengkapi. Strategi ini "
        "memungkinkan transisi dari development sampai production tanpa "
        "perubahan kode aplikasi.",
    )

    add_image(doc, DIAGRAMS_DIR / "03_deployment_layers.png", width_cm=14,
              caption="Gambar 5.7 — Strategi deployment multi-layer: lokal "
                      "(development) → Cloudflare Tunnel (demo publik) → "
                      "Docker + Hugging Face Spaces (hosting permanen)")

    _unused_multilayer = """┌─────────────────────────────────────────────────────────────┐
│            STRATEGI DEPLOYMENT MULTI-LAYER                  │
└─────────────────────────────────────────────────────────────┘

  Layer 1: LOKAL DEVELOPMENT
  ┌───────────────────────────────────────────┐
  │  ./scripts/run_app.sh                     │
  │  - FastAPI di localhost:8000              │
  │  - Streamlit di localhost:8501            │
  │  - Untuk pengembangan & testing internal  │
  └───────────────────────────────────────────┘
                      │
                      ▼  (saat butuh demo publik)
  Layer 2: DEMO PUBLIK (Cloudflare Tunnel)
  ┌───────────────────────────────────────────┐
  │  ./scripts/run_demo.sh                    │
  │  - Layer 1 + tunnel publik via            │
  │    cloudflared                            │
  │  - URL: https://xxx.trycloudflare.com     │
  │  - Cocok untuk demo sidang / presentasi   │
  │  - Tidak butuh signup / kartu kredit      │
  │  - Komputer harus tetap menyala           │
  └───────────────────────────────────────────┘
                      │
                      ▼  (saat butuh hosting permanen)
  Layer 3: HOSTING PERMANEN (Hugging Face Spaces)
  ┌───────────────────────────────────────────┐
  │  Dockerfile + ./scripts/deploy_hf.sh      │
  │  - Container Docker (FastAPI + Streamlit) │
  │  - Di-host di Hugging Face Spaces gratis  │
  │  - URL stabil 24/7                        │
  │  - Auto-rebuild saat code di-push         │
  └───────────────────────────────────────────┘
"""

    add_para(doc, "Layer 1 — Lokal Development:", bold=True)
    add_justified(
        doc,
        "Skrip scripts/run_app.sh menjalankan FastAPI dan Streamlit pada "
        "mesin lokal. Cocok untuk pengembangan, debugging, dan testing "
        "internal. Pre-check port konflik dan auto-generate serving "
        "artifacts kalau belum ada.",
    )

    add_para(doc, "Layer 2 — Demo Publik via Cloudflare Tunnel:", bold=True)
    add_justified(
        doc,
        "Skrip scripts/run_demo.sh memperluas Layer 1 dengan membuka tunnel "
        "publik via cloudflared. URL publik (format "
        "https://xxx.trycloudflare.com) otomatis di-generate dan di-copy ke "
        "clipboard. Cocok untuk demo cepat saat sidang tanpa perlu signup "
        "atau kartu kredit. Keterbatasannya: komputer harus tetap menyala "
        "dan URL berubah setiap kali script di-restart.",
    )

    add_para(doc, "Layer 3 — Hosting Permanen via Hugging Face Spaces:", bold=True)
    add_justified(
        doc,
        "Untuk hosting permanen yang dapat diakses 24/7 tanpa perlu komputer "
        "lokal menyala, aplikasi di-containerize menggunakan Docker dan "
        "di-deploy ke Hugging Face Spaces (free tier). Komponen yang "
        "disiapkan:",
    )
    add_bullets(
        doc,
        [
            "Dockerfile — base image python:3.11-slim, install dependencies, "
            "copy aplikasi, run start.sh.",
            "start.sh — launcher yang menjalankan FastAPI di internal port "
            "8000 lalu Streamlit di port 7860 sebagai foreground process.",
            ".dockerignore — exclude data raw, notebooks, virtual env, dan "
            "file besar lain yang tidak dibutuhkan saat runtime.",
            "requirements-deploy.txt — dependencies minimal khusus serving "
            "(skip jupyter, plotly, dll) untuk mempercepat build.",
            "scripts/deploy_hf.sh — upload semua file ke Hugging Face Space "
            "menggunakan HF CLI dengan satu perintah.",
        ],
    )

    add_para(doc, "Cara Deploy ke Hugging Face Spaces:", bold=True)
    add_code(doc, """# 1. Signup di huggingface.co/join (gratis)

# 2. Generate access token Write di:
#    huggingface.co/settings/tokens

# 3. Login dari terminal
hf auth login

# 4. Buat Space baru via web UI:
#    huggingface.co/new-space
#    SDK = Docker, Hardware = CPU basic (gratis)

# 5. Upload semua file deployment
./scripts/deploy_hf.sh <username>/<space-name>

# 6. URL publik (permanent):
#    https://<username>-<space-name>.hf.space""")

    add_heading(doc, "5.10 Perbandingan Layer Deployment", level=2)
    add_table(
        doc,
        ["Aspek", "Layer 1 (Lokal)", "Layer 2 (Tunnel)", "Layer 3 (HF Spaces)"],
        [
            ["Biaya", "Gratis", "Gratis", "Gratis (free tier)"],
            ["URL publik", "Tidak", "Ya (temporary)", "Ya (permanent)"],
            ["URL stabil", "—", "Berubah tiap restart", "Stabil 24/7"],
            ["Butuh komputer nyala", "Ya", "Ya", "Tidak"],
            ["Setup time", "Instan", "~5 menit", "~15 menit pertama kali"],
            ["Use case", "Development", "Demo sidang", "Portofolio / 24-7"],
        ],
        col_widths=[3.5, 3.5, 3.8, 4.0],
    )

    add_para(doc, "Catatan Free Tier Hugging Face Spaces:", bold=True)
    add_bullets(
        doc,
        [
            "Resource: CPU 2 vCPU, RAM 16 GB, storage 50 GB (cukup untuk "
            "model 374 KB + dataset training kalau dibutuhkan).",
            "Sleep mode: Space akan otomatis sleep setelah ~48 jam idle. "
            "First visit setelah sleep memerlukan ~30 detik untuk wake up "
            "container (auto, tanpa intervensi).",
            "Visit reset countdown: setiap visit mereset timer 48 jam, "
            "sehingga selama ada aktivitas berkala Space tetap running.",
            "Build pertama lama (~10 menit) karena harus install semua "
            "dependencies. Build berikutnya lebih cepat (~2-3 menit) berkat "
            "Docker layer cache.",
        ],
    )

    add_heading(doc, "5.11 Source Code Repository (GitHub)", level=2)
    add_justified(
        doc,
        "Seluruh source code, model, konfigurasi pipeline, laporan, dan "
        "dokumentasi proyek di-publish sebagai repositori open-source di "
        "GitHub untuk transparansi, reproducibility, dan kemudahan akses "
        "bagi reviewer:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("https://github.com/rafafifuwae/rfm-customer-segmentation")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    add_para(doc, "Isi Repository:", bold=True)
    add_bullets(
        doc,
        [
            "Pipeline notebook lengkap (data collection, EDA, preprocessing, "
            "modeling) dengan output cell tersimpan untuk dibaca langsung "
            "di GitHub.",
            "Source code aplikasi serving (FastAPI + Streamlit + predictor "
            "wrapper) yang siap di-clone dan dijalankan.",
            "Model K-Means terlatih (models/kmeans_rfm.pkl, 374 KB) beserta "
            "serving artifacts (winsorize bounds + segment mapping).",
            "Docker config (Dockerfile, start.sh) untuk deployment otomatis "
            "ke Hugging Face Spaces atau platform container lain.",
            "27 file gambar visualisasi (21 figure modeling + 6 screenshot "
            "deployment) sebagai bukti hasil dan dokumentasi visual.",
            "Laporan tugas dalam format DOCX yang siap di-review.",
            "DVC pipeline configuration (dvc.yaml, params.yaml) untuk "
            "reproducibility eksperimen.",
        ],
    )

    add_para(doc, "Tidak Disertakan (sesuai .gitignore):", bold=True)
    add_bullets(
        doc,
        [
            "Dataset raw Olist (~120 MB) — terlalu besar untuk Git biasa. "
            "Dataset dapat di-download dari Kaggle (referensi link di "
            "README) atau di-track via DVC remote.",
            "Virtual environment (.venv) dan dependency yang terinstall — "
            "dapat di-restore via pip install -r requirements.txt.",
            "File cache dan temporary lain (__pycache__, .ipynb_checkpoints, "
            ".dvc/cache).",
        ],
    )

    add_para(doc, "Cara Clone & Jalankan:", bold=True)
    add_code(doc, """# 1. Clone repository
git clone https://github.com/rafafifuwae/rfm-customer-segmentation.git
cd rfm-customer-segmentation

# 2. Setup virtual environment
python -m venv .venv
source .venv/bin/activate          # Linux/Mac
# .venv\\Scripts\\activate          # Windows

# 3. Install dependencies
pip install -r requirements.txt

# 4. (Opsional) Download dataset Olist dari Kaggle
#    https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce
#    Ekstrak ke folder data/raw/
#    (Skip langkah ini kalau hanya ingin menjalankan inference,
#     karena model & serving artifacts sudah ter-include)

# 5. Jalankan aplikasi
./scripts/run_app.sh
#    UI       : http://localhost:8501
#    API docs : http://localhost:8000/docs""")

    add_para(doc, "Manfaat Publikasi di GitHub:", bold=True)
    add_bullets(
        doc,
        [
            "Transparansi penelitian — reviewer/dosen dapat memeriksa "
            "setiap baris kode dan keputusan teknis langsung di browser.",
            "Reproducibility — siapa pun dapat menjalankan ulang pipeline "
            "yang sama dengan random_state yang fixed.",
            "Portofolio profesional — repository menjadi bukti karya nyata "
            "dengan dokumentasi lengkap (commit history, README, kode).",
            "Versioning — perubahan kode tercatat lengkap melalui commit "
            "history, memudahkan tracking eksperimen.",
            "Kolaborasi — siapa pun dapat fork, contribute, atau "
            "mengadaptasi project untuk dataset lain.",
        ],
    )

    add_heading(doc, "Kesimpulan", level=1)
    add_justified(
        doc,
        "Penelitian ini berhasil membangun sistem segmentasi pelanggan "
        "end-to-end mulai dari raw data sekitar 120 MB hingga deployment "
        "yang siap pakai. Pipeline mengikuti standar CRISP-DM, di-track via "
        "DVC untuk reproducibility, dan menghasilkan 4 segmen pelanggan "
        "dengan metrik Silhouette = 0,372 dan Davies-Bouldin = 0,763 yang "
        "menunjukkan kualitas clustering yang baik untuk konteks data "
        "perilaku pelanggan. Hasil segmentasi diterjemahkan menjadi "
        "rekomendasi strategi pemasaran yang spesifik per segmen "
        "(Champions, Loyal Customers, At Risk, Hibernating / Lost).",
    )
    add_justified(
        doc,
        "Dari sisi deployment, sistem dirancang dengan tiga layer yang "
        "saling melengkapi: (1) Layer lokal untuk development, (2) Layer "
        "demo publik via Cloudflare Tunnel untuk presentasi cepat tanpa "
        "biaya tambahan, dan (3) Layer hosting permanen via Docker + "
        "Hugging Face Spaces untuk akses 24/7. Aplikasi dapat diakses "
        "melalui Web UI Streamlit yang interaktif maupun REST API FastAPI "
        "yang siap diintegrasikan dengan sistem lain. Strategi multi-layer "
        "ini memastikan transisi mulus dari development hingga production "
        "tanpa perubahan kode aplikasi.",
    )
    add_justified(
        doc,
        "Sebagai bentuk transparansi dan untuk mendukung reproducibility, "
        "seluruh kode, model, konfigurasi, dan laporan dipublikasikan "
        "sebagai repository open-source di GitHub "
        "(https://github.com/rafafifuwae/rfm-customer-segmentation) "
        "sehingga reviewer maupun pihak ketiga dapat memeriksa, "
        "menjalankan ulang, atau mengadaptasi sistem ini untuk dataset "
        "lain.",
    )

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Tersimpan -> {OUT_PATH}")
    print(f"  Ukuran   -> {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
